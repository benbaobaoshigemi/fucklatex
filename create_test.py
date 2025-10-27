from docx import Document

# 创建测试文档
doc = Document()
doc.add_heading('LaTeX公式测试文档', 0)

doc.add_paragraph('这是一个包含LaTeX公式的测试文档。')
doc.add_paragraph()

doc.add_paragraph('1. 简单的平方公式：$x^2$')
doc.add_paragraph('2. 分数公式：$\\frac{a}{b}$')
doc.add_paragraph('3. 求和公式：$\\sum_{i=1}^{n} i$')
doc.add_paragraph('4. 积分公式：$\\int_{0}^{1} x dx$')
doc.add_paragraph('5. 平方根：$\\sqrt{2}$')
doc.add_paragraph('6. 希腊字母：$\\alpha + \\beta = \\gamma$')
doc.add_paragraph('7. 复杂公式：$E = mc^2$')
doc.add_paragraph()

doc.add_paragraph('段落中混合文本和公式：方程 $ax^2 + bx + c = 0$ 的解为 $x = \\frac{-b \\pm \\sqrt{b^2-4ac}}{2a}$。')

doc.save('test.docx')
print("测试文档已创建：test.docx")