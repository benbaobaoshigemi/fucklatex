# -*- coding: utf-8 -*-
"""
Word LaTeX公式渲染器 - CLI工具
自动将Word文档中的 $...$ LaTeX公式转换为Word公式对象
支持自动渲染为专业格式
"""

import re
import os
import sys
import subprocess
import argparse
import io
import time
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def check_dependencies(auto_install=False):
    """检查并安装必要的依赖"""
    required_packages = ['python-docx']
    missing_packages = []

    print("\n🔍 检查依赖...")
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"   ✅ {package}")
        except ImportError:
            missing_packages.append(package)
            print(f"   ❌ {package} - 未安装")

    if missing_packages:
        print(f"\n⚠️  发现 {len(missing_packages)} 个缺失的依赖包:")
        for pkg in missing_packages:
            print(f"   • {pkg}")

        # 如果启用了自动安装，或者在非交互式环境中
        if auto_install or not sys.stdin.isatty():
            print("\n📦 正在自动安装依赖...")
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install'] + missing_packages)
                print("✅ 依赖安装完成！")
                return True
            except subprocess.CalledProcessError as e:
                print(f"❌ 依赖安装失败: {e}")
                print("请手动运行: pip install " + ' '.join(missing_packages))
                return False
        else:
            # 交互式模式
            try:
                choice = input("\n是否自动安装缺失的依赖？(y/n): ").strip().lower()
                if choice == 'y':
                    print("\n📦 正在安装依赖...")
                    try:
                        subprocess.check_call([sys.executable, '-m', 'pip', 'install'] + missing_packages)
                        print("✅ 依赖安装完成！")
                        return True
                    except subprocess.CalledProcessError as e:
                        print(f"❌ 依赖安装失败: {e}")
                        print("请手动运行: pip install " + ' '.join(missing_packages))
                        return False
                else:
                    print("❌ 请先安装依赖后重新运行程序")
                    return False
            except EOFError:
                print("\n⚠️  非交互式环境，跳过依赖安装")
                print("请手动运行: pip install " + ' '.join(missing_packages))
                return False

    print("✅ 所有依赖已就绪！")
    return True

def check_dollar_signs(file_path):
    """检查$符号是否成对出现"""
    try:
        doc = Document(file_path)
        total_dollar_count = 0

        for para in doc.paragraphs:
            text = para.text
            dollar_count = text.count('$')
            total_dollar_count += dollar_count

        if total_dollar_count % 2 != 0:
            print(f"\n❌ 错误：文档中的$符号数量为奇数 ({total_dollar_count})，无法正确匹配公式对！")
            print("请检查文档中的公式格式，确保每个公式都被一对$符号包围。")
            print("\n💡 提示：")
            print("   • 正确格式：$x^2$")
            print("   • 错误格式：$x^2 或 x^2$")
            return False

        print(f"✅ $符号检查通过 (共 {total_dollar_count} 个，{total_dollar_count//2} 对公式)")
        return True

    except Exception as e:
        print(f"❌ 检查$符号时出错: {e}")
        return False

def get_save_mode(input_path):
    """获取保存模式"""
    print(f"\n💾 请选择保存模式:")
    print(f"   0 - 覆盖原文件 (⚠️ 会替换 {input_path})")
    print(f"   1 - 保存到当前目录 ({os.getcwd()})")
    print(f"   2 - 指定保存路径")

    while True:
        choice = input("\n请选择 (0/1/2): ").strip()

        if choice == '0':
            return input_path  # 覆盖原文件

        elif choice == '1':
            # 保存到当前目录
            base_name = os.path.basename(input_path)
            name, ext = os.path.splitext(base_name)
            output_path = os.path.join(os.getcwd(), f"{name}_processed{ext}")
            return output_path

        elif choice == '2':
            # 自定义路径
            output_path = input("请输入完整的保存路径: ").strip().strip('"')
            if not output_path:
                print("❌ 保存路径不能为空")
                continue

            # 确保有扩展名
            if not output_path.lower().endswith('.docx'):
                output_path += '.docx'

            return output_path

        else:
            print("❌ 无效选择，请输入 0、1 或 2")

def latex_to_unicodemath(latex_str):
    """将LaTeX转换为Word的UnicodeMath格式"""
    replacements = {
        r'\\frac\{([^}]+)\}\{([^}]+)\}': r'(\1)/(\2)',
        r'\\sqrt\{([^}]+)\}': r'√(\1)',
        r'\\sum': '∑',
        r'\\int': '∫',
        r'\\prod': '∏',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\delta': 'δ',
        r'\\epsilon': 'ε',
        r'\\pi': 'π',
        r'\\theta': 'θ',
        r'\\lambda': 'λ',
        r'\\mu': 'μ',
        r'\\sigma': 'σ',
        r'\\omega': 'ω',
        r'\\Gamma': 'Γ',
        r'\\Delta': 'Δ',
        r'\\Theta': 'Θ',
        r'\\Lambda': 'Λ',
        r'\\Sigma': 'Σ',
        r'\\Omega': 'Ω',
        r'\\pm': '±',
        r'\\times': '×',
        r'\\div': '÷',
        r'\\le': '≤',
        r'\\ge': '≥',
        r'\\ne': '≠',
        r'\\approx': '≈',
        r'\\infty': '∞',
        r'\\rightarrow': '→',
        r'\\leftarrow': '←',
        r'\\Rightarrow': '⇒',
        r'\\Leftarrow': '⇐',
    }

    result = latex_str
    for pattern, replacement in replacements.items():
        result = re.sub(pattern, replacement, result)

    return result

def create_omml_formula(latex_str):
    """创建OMML公式元素"""
    try:
        unicodemath = latex_to_unicodemath(latex_str)
        unicodemath_escaped = unicodemath.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        omml = f'''<m:oMath {nsdecls("m", "w")}>
    <m:r>
        <w:rPr>
            <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:cs="Cambria Math"/>
        </w:rPr>
        <m:t>{unicodemath_escaped}</m:t>
    </m:r>
</m:oMath>'''

        return omml, unicodemath
    except Exception as e:
        return None, None

def auto_render_formulas(file_path):
    """自动打开Word并将公式渲染为专业格式"""
    
    print("\n" + "="*70)
    print("🎨 自动渲染公式为专业格式")
    print("="*70)
    
    # 检查pywin32是否已安装
    try:
        import win32com.client
    except ImportError:
        print("⚠️  缺少 pywin32 依赖")
        try:
            choice = input("是否自动安装 pywin32？(y/n): ").strip().lower()
            if choice == 'y':
                print("\n📦 正在安装 pywin32...")
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pywin32'])
                print("✅ pywin32 安装完成")
                import win32com.client
            else:
                print("⏭️  跳过自动渲染")
                return False
        except:
            print("❌ 安装失败，跳过自动渲染")
            return False
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"❌ 错误：文件不存在: {file_path}")
        return False
    
    # 获取绝对路径
    abs_path = os.path.abspath(file_path)
    print(f"\n📂 文档路径: {abs_path}")
    
    word = None
    doc = None
    
    try:
        print("\n⏳ 正在启动Word...")
        # 创建Word应用程序实例
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True  # 显示Word窗口
        
        print("⏳ 正在打开文档...")
        # 打开文档
        doc = word.Documents.Open(abs_path)
        
        print("✅ 文档已打开")
        print("\n🔍 正在扫描公式...")
        
        # 等待文档完全加载
        time.sleep(2)
        
        # 获取文档范围内的所有公式
        formula_count = 0
        all_omaths = []
        
        # 遍历所有故事范围（包括主文档、页眉页脚等）
        for story in doc.StoryRanges:
            while story:
                # 获取当前故事范围中的公式
                if story.OMaths.Count > 0:
                    for i in range(1, story.OMaths.Count + 1):
                        all_omaths.append(story.OMaths.Item(i))
                        formula_count += 1
                
                # 移动到下一个故事范围
                try:
                    story = story.NextStoryRange
                except:
                    break
        
        print(f"📊 检测到 {formula_count} 个公式对象")
        
        if formula_count == 0:
            print("⚠️  警告：文档中没有检测到公式对象")
            return False
        else:
            print("\n🔄 正在将公式从线性格式转换为专业格式...")
            print("💡 说明：BuildUp() 会将 UnicodeMath 线性文本构建为二维数学公式")
            
            # 遍历所有公式并转换为专业格式
            converted_count = 0
            failed_count = 0
            
            for i, omath in enumerate(all_omaths, 1):
                try:
                    # 使用BuildUp()方法将线性格式转换为专业格式
                    omath.BuildUp()
                    converted_count += 1
                    
                    # 每10个公式显示一次进度
                    if converted_count % 10 == 0:
                        print(f"   ⏳ 已处理 {converted_count}/{formula_count} 个公式...")
                    
                except Exception as e:
                    failed_count += 1
                    if failed_count <= 3:  # 只显示前3个错误
                        print(f"   ⚠️  公式 {i} 转换失败: {e}")
            
            print(f"\n✅ 渲染完成！")
            print(f"📊 统计：")
            print(f"   • 总计: {formula_count} 个公式")
            print(f"   • 成功: {converted_count} 个")
            if failed_count > 0:
                print(f"   • 失败: {failed_count} 个")
            
            # 保存文档
            print("\n💾 正在保存文档...")
            doc.Save()
            print("✅ 文档已保存")
        
        print("\n" + "="*70)
        print("✅ 自动渲染完成")
        print("="*70)
        print("\n💡 Word将保持打开状态，您可以查看结果")
        
        return True
        
    except Exception as e:
        print(f"\n❌ 渲染错误: {e}")
        import traceback
        traceback.print_exc()
        return False

def process_document(input_path, output_path):
    """处理Word文档"""
    try:
        print(f"\n{'='*70}")
        print(f"� 开始处理文档")
        print(f"{'='*70}")
        print(f"📂 输入: {input_path}")
        print(f"📄 输出: {output_path}")

        doc = Document(input_path)
        processed_count = 0
        failed_count = 0

        print(f"\n🔍 扫描文档...")

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            matches = list(re.finditer(r'\$(.*?)\$', text))

            if matches:
                print(f"📝 段落 {para_idx + 1}: {len(matches)} 个公式")

                # 清空段落
                p_element = para._p
                p_element.clear_content()

                # 重建段落
                last_end = 0
                for match in matches:
                    latex_str = match.group(1)
                    start, end = match.span()

                    # 添加公式前的文本
                    if start > last_end:
                        para.add_run(text[last_end:start])

                    # 创建公式
                    omml, unicodemath = create_omml_formula(latex_str)

                    if omml:
                        try:
                            omml_element = parse_xml(omml)
                            para._p.append(omml_element)
                            processed_count += 1
                        except Exception as e:
                            para.add_run(match.group(0))
                            failed_count += 1
                    else:
                        para.add_run(match.group(0))
                        failed_count += 1

                    last_end = end

                # 添加剩余文本
                if last_end < len(text):
                    para.add_run(text[last_end:])

        # 保存文档
        print(f"\n💾 保存文档...")
        doc.save(output_path)

        # 输出结果
        print(f"\n{'='*70}")
        print(f"✅ 处理完成！")
        print(f"{'='*70}")
        print(f"📊 统计:")
        print(f"   • 成功: {processed_count} 个公式")
        if failed_count > 0:
            print(f"   • 失败: {failed_count} 个公式")
        print(f"📄 输出文件: {output_path}")
        print(f"{'='*70}")

        return True

    except Exception as e:
        print(f"\n❌ 处理失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    # 设置环境变量确保UTF-8编码
    os.environ['PYTHONIOENCODING'] = 'utf-8'

    # 在Windows上强制设置控制台编码
    if sys.platform == 'win32':
        try:
            import ctypes
            kernel32 = ctypes.windll.kernel32
            # 设置控制台输出代码页为UTF-8
            kernel32.SetConsoleOutputCP(65001)
            kernel32.SetConsoleCP(65001)
        except:
            pass
    parser = argparse.ArgumentParser(
        description='Word LaTeX公式渲染器 - 将Word文档中的 $...$ LaTeX公式转换为Word公式对象',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python main.py                           # 交互模式
  python main.py document.docx             # 处理指定文档（保存到当前目录）
  python main.py document.docx -o output.docx  # 指定输出文件
  python main.py document.docx --overwrite     # 覆盖原文件

保存模式:
  无指定: 保存到当前目录 (filename_processed.docx)
  -o output.docx: 指定输出文件路径
  --overwrite: 覆盖原文件
        """
    )

    parser.add_argument('input_file', nargs='?', help='输入的Word文档路径 (.docx)')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--overwrite', action='store_true', help='覆盖原文件')
    parser.add_argument('--auto-install', action='store_true', help='自动安装缺失的依赖包')

    args = parser.parse_args()

    print("="*70)
    print(" Word LaTeX公式渲染器 - CLI工具")
    print("="*70)
    print("功能：自动将Word文档中的 $...$ LaTeX公式转换为Word公式对象")
    print("版本：2.0 - CLI版")
    print("="*70)

    # 1. 检查依赖
    if not check_dependencies(args.auto_install):
        return

    # 如果提供了命令行参数，使用参数模式
    if args.input_file:
        file_path = args.input_file

        # 验证输入文件
        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return

        if not file_path.lower().endswith('.docx'):
            print("❌ 只支持 .docx 格式的Word文档")
            return

        # 确定输出路径
        if args.overwrite:
            output_path = file_path
        elif args.output:
            output_path = args.output
            if not output_path.lower().endswith('.docx'):
                output_path += '.docx'
        else:
            # 默认保存到当前目录
            base_name = os.path.basename(file_path)
            name, ext = os.path.splitext(base_name)
            output_path = os.path.join(os.getcwd(), f"{name}_processed{ext}")

    else:
        # 交互模式
        # 2. 获取输入文件
        while True:
            file_path = input("\n📂 请输入Word文档路径: ").strip().strip('"')
            if not file_path:
                print("❌ 文件路径不能为空")
                continue

            if not os.path.exists(file_path):
                print(f"❌ 文件不存在: {file_path}")
                continue

            if not file_path.lower().endswith('.docx'):
                print("❌ 只支持 .docx 格式的Word文档")
                continue

            break

        # 3. 检查$符号
        if not check_dollar_signs(file_path):
            return

        # 4. 选择保存模式
        output_path = get_save_mode(file_path)

    # 3/5. 检查$符号（参数模式也需要检查）
    if not check_dollar_signs(file_path):
        return

    # 6. 确认操作（仅交互模式）
    if not args.input_file:
        print(f"\n🔄 准备执行:")
        print(f"   输入文件: {file_path}")
        print(f"   输出文件: {output_path}")

        if output_path == file_path:
            confirm = input("\n⚠️  将覆盖原文件，确定继续？(y/n): ").strip().lower()
            if confirm != 'y':
                print("❌ 操作已取消")
                return

    # 7. 执行处理
    success = process_document(file_path, output_path)

    if success:
        print(f"\n🎉 处理成功！")
        print(f"💡 提示: 已生成 {output_path}")
        
        # 询问是否自动渲染
        print("\n" + "="*70)
        print("🎨 自动渲染选项")
        print("="*70)
        print("现在可以自动打开Word并将公式转换为专业的二维格式")
        print("(将线性格式如 '(a)/(b)' 渲染为漂亮的分数显示)")
        
        # 在命令行模式下，也询问是否渲染
        render_choice = input("\n是否自动渲染公式？(y/n): ").strip().lower()
        
        if render_choice == 'y':
            render_success = auto_render_formulas(output_path)
            
            if render_success:
                print("\n✅ 完整流程已完成！")
                print("📂 Word文档已打开，您可以查看渲染后的公式效果")
            else:
                print("\n⚠️  自动渲染未成功，但文档已生成")
                print("💡 您可以手动打开文档并在Word中操作")
        else:
            print("\n⏭️  跳过自动渲染")
            print("💡 您可以稍后手动打开文档")
            
            # 询问是否打开文件
            open_choice = input("\n是否现在打开文档？(y/n): ").strip().lower()
            if open_choice == 'y':
                try:
                    os.startfile(output_path)
                    print(f"✅ 已打开文档")
                except:
                    print(f"❌ 无法自动打开，请手动打开: {output_path}")
    else:
        print(f"\n❌ 处理失败！")

    # 仅在交互模式下等待退出
    if not args.input_file:
        input("\n按Enter键退出...")

if __name__ == "__main__":
    main()
