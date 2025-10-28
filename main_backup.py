# -*- coding: utf-8 -*-
"""
Word LaTeX公式渲染器 - CLI工具
自动将Word文档中的 $...$ LaTeX公式转换为Word公式对象
支持自动渲染为专业格式
支持LaTeX拼写检查（检测命令中的空格插入等错误）

核心原理：
1. 保留 $LaTeX$ 格式不变
2. 用Word COM API模拟Alt+=操作，将LaTeX标记为公式对象
3. 调用BuildUp()方法渲染为专业的二维格式

优势：利用Word原生LaTeX支持，无需维护命令映射表
"""

import re
import os
import sys
import subprocess
import argparse
import io
import time

def print_banner():
    """显示炫酷的启动横幅"""
    banner = r"""
╔════════════════════════════════════════════════════════════════════╗
║                                                                    ║
║   ██╗    ██╗ ██████╗ ██████╗ ██████╗     ██╗      █████╗ ████████╗║
║   ██║    ██║██╔═══██╗██╔══██╗██╔══██╗    ██║     ██╔══██╗╚══██╔══╝║
║   ██║ █╗ ██║██║   ██║██████╔╝██║  ██║    ██║     ███████║   ██║   ║
║   ██║███╗██║██║   ██║██╔══██╗██║  ██║    ██║     ██╔══██║   ██║   ║
║   ╚███╔███╔╝╚██████╔╝██║  ██║██████╔╝    ███████╗██║  ██║   ██║   ║
║    ╚══╝╚══╝  ╚═════╝ ╚═╝  ╚═╝╚═════╝     ╚══════╝╚═╝  ╚═╝   ╚═╝   ║
║                                                                    ║
║              ███████╗ ██████╗ ██████╗ ███╗   ███╗██╗   ██╗██╗      ║
║              ██╔════╝██╔═══██╗██╔══██╗████╗ ████║██║   ██║██║      ║
║              █████╗  ██║   ██║██████╔╝██╔████╔██║██║   ██║██║      ║
║              ██╔══╝  ██║   ██║██╔══██╗██║╚██╔╝██║██║   ██║██║      ║
║              ██║     ╚██████╔╝██║  ██║██║ ╚═╝ ██║╚██████╔╝███████╗ ║
║              ╚═╝      ╚═════╝ ╚═╝  ╚═╝╚═╝     ╚═╝ ╚═════╝ ╚══════╝ ║
║                                                                    ║
║              ██████╗ ███████╗███╗   ██╗██████╗ ███████╗██████╗     ║
║              ██╔══██╗██╔════╝████╗  ██║██╔══██╗██╔════╝██╔══██╗    ║
║              ██████╔╝█████╗  ██╔██╗ ██║██║  ██║█████╗  ██████╔╝    ║
║              ██╔══██╗██╔══╝  ██║╚██╗██║██║  ██║██╔══╝  ██╔══██╗    ║
║              ██║  ██║███████╗██║ ╚████║██████╔╝███████╗██║  ██║    ║
║              ╚═╝  ╚═╝╚══════╝╚═╝  ╚═══╝╚═════╝ ╚══════╝╚═╝  ╚═╝    ║
║                                                                    ║
╚════════════════════════════════════════════════════════════════════╝

    🚀 LaTeX → Word 公式转换器 & 自动渲染工具
    📦 Version 3.0.0 - Word Native API
    ⚡ 一键转换 | 自动渲染 | 专业格式
    
    💡 功能：将 $...$ LaTeX 公式转换为 Word 原生公式对象
    🎨 特色：利用Word原生LaTeX支持，100%命令兼容
    
"""
    print(banner)

def check_dependencies(auto_install=False):
    """检查并安装必要的依赖"""
    print("🔍 检查依赖...")
    
    # 尝试导入 docx
    try:
        import docx
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        print("   ✅ python-docx")
        return True
    except ImportError:
        print("   ❌ python-docx - 未安装")
        
        # 如果启用了自动安装，或者在非交互式环境中
        if auto_install or not sys.stdin.isatty():
            print("\n📦 正在自动安装依赖...")
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
                print("✅ 依赖安装完成！")
                return True
            except subprocess.CalledProcessError as e:
                print(f"❌ 依赖安装失败: {e}")
                print("请手动运行: pip install python-docx")
                return False
        else:
            # 交互式模式
            try:
                choice = input("\n是否自动安装 python-docx？(y/n): ").strip().lower()
                if choice == 'y':
                    print("\n📦 正在安装依赖...")
                    try:
                        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
                        print("✅ 依赖安装完成！")
                        return True
                    except subprocess.CalledProcessError as e:
                        print(f"❌ 依赖安装失败: {e}")
                        print("请手动运行: pip install python-docx")
                        return False
                else:
                    print("❌ 请先安装依赖后重新运行程序")
                    return False
            except EOFError:
                print("\n⚠️  非交互式环境，跳过依赖安装")
                print("请手动运行: pip install python-docx")
                return False

def check_dollar_signs(file_path):
    """检查$符号是否成对出现"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        total_dollar_count = 0

        for para in doc.paragraphs:
            text = para.text
            dollar_count = text.count('$')
            total_dollar_count += dollar_count

        if total_dollar_count % 2 != 0:
            print(f"\n❌ 错误：文档中的$符号数量为奇数 ({total_dollar_count})，无法正确匹配公式对！")
            print("请检查文档中的公式格式，确保每个公式都被一对$符号包围。")
            print("\n💡 提示：")
            print("   • 正确格式：$x^2$")
            print("   • 错误格式：$x^2 或 x^2$")
            return False

        print(f"✅ $符号检查通过 (共 {total_dollar_count} 个，{total_dollar_count//2} 对公式)")
        return True

    except Exception as e:
        print(f"❌ 检查$符号时出错: {e}")
        return False

def get_save_mode(input_path):
    """获取保存模式"""
    print(f"\n💾 请选择保存模式:")
    print(f"   0 - 覆盖原文件 (⚠️ 会替换 {input_path})")
    print(f"   1 - 保存到当前目录 ({os.getcwd()})")
    print(f"   2 - 指定保存路径")

    while True:
        choice = input("\n请选择 (0/1/2): ").strip()

        if choice == '0':
            return input_path  # 覆盖原文件

        elif choice == '1':
            # 保存到当前目录
            base_name = os.path.basename(input_path)
            name, ext = os.path.splitext(base_name)
            output_path = os.path.join(os.getcwd(), f"{name}_processed{ext}")
            return output_path

        elif choice == '2':
            # 自定义路径
            output_path = input("请输入完整的保存路径: ").strip().strip('"')
            if not output_path:
                print("❌ 保存路径不能为空")
                continue

            # 确保有扩展名
            if not output_path.lower().endswith('.docx'):
                output_path += '.docx'

            return output_path

        else:
            print("❌ 无效选择，请输入 0、1 或 2")

# ============================================================================
# ⚠️  以下函数已废弃（保留仅供参考）
# ============================================================================
# 原因：Word原生支持LaTeX转换功能（Alt+=）
# 新策略：直接保留 $LaTeX$ 格式，然后用Word COM API模拟Alt+=操作
# 优势：
#   1. 无需维护命令映射表（Word支持完整的LaTeX语法）
#   2. 代码简化200+行
#   3. 完美兼容性（Word官方实现）
# ============================================================================

def latex_to_unicodemath(latex_str):
    """
    ⚠️  已废弃：将LaTeX转换为Word的UnicodeMath格式
    
    废弃原因：Word原生支持LaTeX→公式转换（Alt+=）
    新方案：见 auto_render_formulas() 函数
    
    改进策略（已不再使用）:
    1. 自动按命令长度排序（长命令优先匹配）
    2. 使用单词边界 \\b 防止误匹配
    3. 避免 \\le 匹配 \\left、\\leftarrow 等
    """
    # 定义所有替换规则
    replacements = {
        # 带参数的复杂命令
        r'\\frac\{([^}]+)\}\{([^}]+)\}': r'(\1)/(\2)',
        r'\\sqrt\{([^}]+)\}': r'√(\1)',
        
        # 🆕 物理学常用符号（必须在其他命令之前，避免被误匹配）
        r'\\dagger': '†',     # 产生算符/厄米共轭
        r'\\hbar': 'ℏ',       # 约化普朗克常数
        r'\\partial': '∂',    # 偏微分
        r'\\nabla': '∇',      # 梯度/纳布拉算符
        r'\\cdot': '·',       # 点乘
        r'\\circ': '∘',       # 圆圈/复合函数
        
        # 所有简单命令（会自动按长度排序）
        r'\\leftarrow': '←',
        r'\\Leftarrow': '⇐',
        r'\\rightarrow': '→',
        r'\\Rightarrow': '⇒',
        r'\\epsilon': 'ε',
        r'\\lambda': 'λ',
        r'\\Lambda': 'Λ',
        r'\\approx': '≈',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\Gamma': 'Γ',
        r'\\delta': 'δ',
        r'\\Delta': 'Δ',
        r'\\theta': 'θ',
        r'\\Theta': 'Θ',
        r'\\sigma': 'σ',
        r'\\Sigma': 'Σ',
        r'\\omega': 'ω',
        r'\\Omega': 'Ω',
        r'\\infty': '∞',
        r'\\times': '×',
        r'\\prod': '∏',
        r'\\sum': '∑',
        r'\\int': '∫',
        r'\\div': '÷',
        r'\\mu': 'μ',
        r'\\pi': 'π',
        r'\\pm': '±',
        r'\\le': '≤',  # 会在 \leftarrow 之后处理
        r'\\ge': '≥',
        r'\\ne': '≠',
    }
    
    # � 关键改进：将替换规则按命令长度排序（长的优先）
    # 这样 \leftarrow 会在 \le 之前处理，避免误匹配
    def get_pattern_length(pattern):
        """提取正则表达式中实际命令的长度"""
        # 移除 \\, \{, \}, [, ], +, (, ) 等正则符号
        import re
        clean = re.sub(r'\\\\|\\{|\\}|\[|\]|\+|\(|\)|\^', '', pattern)
        return len(clean)
    
    sorted_replacements = sorted(
        replacements.items(),
        key=lambda x: get_pattern_length(x[0]),
        reverse=True  # 长度降序
    )
    
    result = latex_str
    for pattern, replacement in sorted_replacements:
        # 对于简单命令（不含 { } 的），添加单词边界
        if '{' not in pattern and '}' not in pattern:
            pattern = pattern + r'\b'
        result = re.sub(pattern, replacement, result)
    
    # 🆕 移除单个符号周围的多余花括号
    # {†} → †, {ℏ} → ℏ, {∂} → ∂ 等
    result = re.sub(r'\{([†ℏ∂∇·∘α-ωΑ-Ω∑∫∏±×÷≤≥≠≈∞←→⇐⇒√])\}', r'\1', result)
    
    return result

def create_omml_formula(latex_str):
    """
    ⚠️  已废弃：创建OMML公式元素
    
    废弃原因：不再需要手动构建OMML XML
    新方案：Word COM API 的 OMaths.Add() 方法自动处理
    """
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        unicodemath = latex_to_unicodemath(latex_str)
        unicodemath_escaped = unicodemath.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        omml = f'''<m:oMath {nsdecls("m", "w")}>
    <m:r>
        <w:rPr>
            <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:cs="Cambria Math"/>
        </w:rPr>
        <m:t>{unicodemath_escaped}</m:t>
    </m:r>
</m:oMath>'''

        return omml, unicodemath
    except Exception as e:
        return None, None

def check_and_close_word_document(file_path):
    """检查Word文档是否被打开，如果打开则询问是否关闭"""
    try:
        import win32com.client
        from win32com.client import GetObject
        import pythoncom
        
        abs_path = os.path.abspath(file_path)
        
        try:
            # 尝试获取现有的Word应用程序实例
            word = win32com.client.GetActiveObject("Word.Application")
            
            # 检查文档是否已打开
            for doc in word.Documents:
                if os.path.abspath(doc.FullName).lower() == abs_path.lower():
                    print(f"\n⚠️  检测到文档已在Word中打开")
                    print(f"📂 文档: {os.path.basename(file_path)}")
                    
                    choice = input("\n是否保存并关闭文档以继续处理？(y/n): ").strip().lower()
                    
                    if choice == 'y':
                        try:
                            # 保存文档
                            if doc.Saved == False:
                                print("💾 正在保存文档...")
                                doc.Save()
                                print("✅ 文档已保存")
                            
                            # 关闭文档
                            print("🔒 正在关闭文档...")
                            doc.Close()
                            print("✅ 文档已关闭")
                            time.sleep(1)  # 等待文档完全关闭
                            return True
                        except Exception as e:
                            print(f"❌ 关闭文档失败: {e}")
                            return False
                    else:
                        print("❌ 用户取消操作")
                        return False
        except:
            # 没有运行的Word实例或文档未打开
            pass
        
        return True
        
    except ImportError:
        # pywin32未安装，跳过检查
        return True
    except Exception as e:
        # 其他错误，继续执行
        return True

def auto_render_formulas(file_path):
    """
    自动打开Word并将公式渲染为专业格式
    
    新流程（基于Word原生API）：
    1. 打开Word文档
    2. 查找所有 $...$ LaTeX公式
    3. 调用 OMaths.Add() 将其转换为公式对象（等价于 Alt+=）
    4. 调用 BuildUp() 渲染为专业格式
    
    参考实现：
    Microsoft官方VBA示例（微软社区回复）：
    ```vba
    Sub ConvertSelectionToEquation() 
        With doc.Content.Find 
            .Text = "$$*$$" 
            .MatchWildcards = True 
            Do While .Execute 
                Set eq = Selection.Range.OMaths.Add(Selection.Range) 
            Loop 
        End With 
        For i = 1 To doc.OMaths.Count 
            doc.OMaths(i).BuildUp 
        Next i 
    End Sub
    ```
    本函数是上述VBA代码的Python COM API等价实现。
    """
    
    print("\n" + "="*70)
    print("🎨 自动转换并渲染公式")
    print("="*70)
    
    # 检查pywin32是否已安装
    try:
        import win32com.client
    except ImportError:
        print("⚠️  缺少 pywin32 依赖")
        try:
            choice = input("是否自动安装 pywin32？(y/n): ").strip().lower()
            if choice == 'y':
                print("\n📦 正在安装 pywin32...")
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pywin32'])
                print("✅ pywin32 安装完成")
                import win32com.client
            else:
                print("⏭️  跳过自动渲染")
                return False
        except:
            print("❌ 安装失败，跳过自动渲染")
            return False
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"❌ 错误：文件不存在: {file_path}")
        return False
    
    # 获取绝对路径
    abs_path = os.path.abspath(file_path)
    print(f"\n📂 文档路径: {abs_path}")
    
    word = None
    doc = None
    
    try:
        print("\n⏳ 正在启动Word...")
        # 创建Word应用程序实例
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True  # 显示Word窗口
        
        print("⏳ 正在打开文档...")
        # 打开文档
        doc = word.Documents.Open(abs_path)
        
        print("✅ 文档已打开")
        print("\n🔍 第一步：查找并转换 $LaTeX$ 为公式对象...")
        print("💡 说明：模拟Word的Alt+=功能，将LaTeX文本标记为公式")
        
        # 等待文档完全加载
        time.sleep(2)
        
        # 第一步：将所有 $...$ 转换为公式对象
        converted_count = 0
        failed_count = 0
        
        # 使用Find对象查找所有 $...$ 模式（参考VBA实现）
        find = doc.Content.Find
        find.ClearFormatting()
        find.Text = r"\$*\$"  # 单个 $ 符号的通配符模式
        find.MatchWildcards = True
        
        # 参考VBA的Do While .Execute循环
        while find.Execute():
            try:
                # 获取找到的Range（参考VBA: Set rng = .Parent）
                found_range = find.Parent
                
                # 在这个Range上调用OMaths.Add()
                # 参考VBA: Selection.Range.OMaths.Add(Selection.Range)
                found_range.OMaths.Add(found_range)
                converted_count += 1
                
                if converted_count % 10 == 0:
                    print(f"   ⏳ 已转换 {converted_count} 个公式...")
                
            except Exception as e:
                failed_count += 1
                if failed_count <= 3:
                    print(f"   ⚠️  转换失败: {e}")
                # VBA中用On Error Resume Next继续
                continue
        
        print(f"\n✅ 第一步完成：{converted_count} 个LaTeX公式已转换为公式对象")
        
        if failed_count > 0:
            print(f"⚠️  {failed_count} 个公式转换失败")
        
        # 第二步：统一渲染所有公式（参考VBA: For i = 1 To doc.OMaths.Count）
        print("\n🔍 第二步：渲染公式为专业格式...")
        print("💡 说明：BuildUp() 会将线性格式构建为二维数学公式")
        
        # 等待一下让Word处理完第一步
        time.sleep(1)
        
        # 直接遍历文档中的所有公式对象（像VBA那样）
        formula_count = doc.OMaths.Count
        print(f"📊 检测到 {formula_count} 个公式对象")
        
        if formula_count == 0:
            print("⚠️  警告：文档中没有检测到公式对象")
            print("� 可能原因：LaTeX语法错误，Word无法识别")
            return False
        else:
            # 参考VBA: For i = 1 To doc.OMaths.Count
            rendered_count = 0
            render_failed = 0
            
            for i in range(1, formula_count + 1):
                try:
                    # 使用BuildUp()方法将线性格式转换为专业格式
                    # 对应VBA: doc.OMaths(i).BuildUp
                    doc.OMaths.Item(i).BuildUp()
                    rendered_count += 1
                    
                    # 每10个公式显示一次进度
                    if rendered_count % 10 == 0:
                        print(f"   ⏳ 已渲染 {rendered_count}/{formula_count} 个公式...")
                    
                except Exception as e:
                    render_failed += 1
                    if render_failed <= 3:  # 只显示前3个错误
                        print(f"   ⚠️  公式 {i} 渲染失败: {e}")
            
            print(f"\n✅ 第二步完成！")
            print(f"📊 最终统计：")
            print(f"   • LaTeX→公式对象: {converted_count} 个")
            print(f"   • 渲染为专业格式: {rendered_count} 个")
            if failed_count > 0:
                print(f"   • 转换失败: {failed_count} 个")
            if render_failed > 0:
                print(f"   • 渲染失败: {render_failed} 个")
            
            # 保存文档
            print("\n💾 正在保存文档...")
            doc.Save()
            print("✅ 文档已保存")
        
        print("\n" + "="*70)
        print("✅ 自动转换与渲染完成")
        print("="*70)
        print("\n💡 Word将保持打开状态，您可以查看结果")
        
        return True
        
    except Exception as e:
        print(f"\n❌ 处理错误: {e}")
        import traceback
        traceback.print_exc()
        return False

def process_document(input_path, output_path):
    """
    处理Word文档 - 使用Word原生LaTeX转换功能
    
    新策略：
    1. 保持 $LaTeX$ 原样不变（不做任何转换）
    2. 保存文档
    3. 用Word COM API批量选中并模拟Alt+=操作
    """
    try:
        from docx import Document
        
        print(f"\n{'='*70}")
        print(f"🔄 开始处理文档")
        print(f"{'='*70}")
        print(f"📂 输入: {input_path}")
        print(f"📄 输出: {output_path}")

        doc = Document(input_path)
        formula_count = 0

        print(f"\n🔍 扫描文档中的LaTeX公式...")

        # 统计公式数量（不做任何修改）
        for para in doc.paragraphs:
            text = para.text
            matches = list(re.finditer(r'\$(.*?)\$', text))
            if matches:
                formula_count += len(matches)

        print(f"📊 检测到 {formula_count} 个LaTeX公式（保持 $...$ 格式）")

        # 直接保存文档（不做任何修改）
        print(f"\n💾 保存文档...")
        doc.save(output_path)

        print(f"\n{'='*70}")
        print(f"✅ 文档准备完成！")
        print(f"{'='*70}")
        print(f"📊 统计:")
        print(f"   • 检测到: {formula_count} 个公式")
        print(f"📄 输出文件: {output_path}")
        print(f"\n� 接下来将使用Word自动转换功能（模拟Alt+=）")
        print(f"{'='*70}")

        return True

    except Exception as e:
        print(f"\n❌ 处理失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    # 设置环境变量确保UTF-8编码
    os.environ['PYTHONIOENCODING'] = 'utf-8'

    # 在Windows上强制设置控制台编码
    if sys.platform == 'win32':
        try:
            import ctypes
            kernel32 = ctypes.windll.kernel32
            # 设置控制台输出代码页为UTF-8
            kernel32.SetConsoleOutputCP(65001)
            kernel32.SetConsoleCP(65001)
        except:
            pass
    
    # 显示炫酷的启动横幅
    print_banner()
    
    parser = argparse.ArgumentParser(
        description='Word LaTeX公式渲染器 - 将Word文档中的 $...$ LaTeX公式转换为Word公式对象',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python main.py                           # 交互模式
  python main.py document.docx             # 处理指定文档（保存到当前目录）
  python main.py document.docx -o output.docx  # 指定输出文件
  python main.py document.docx --overwrite     # 覆盖原文件

保存模式:
  无指定: 保存到当前目录 (filename_processed.docx)
  -o output.docx: 指定输出文件路径
  --overwrite: 覆盖原文件
        """
    )

    parser.add_argument('input_file', nargs='?', help='输入的Word文档路径 (.docx)')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--overwrite', action='store_true', help='覆盖原文件')
    parser.add_argument('--auto-install', action='store_true', help='自动安装缺失的依赖包')

    args = parser.parse_args()

    # 1. 检查依赖
    print("\n" + "="*70)
    if not check_dependencies(args.auto_install):
        return
    print("="*70)

    # 如果提供了命令行参数，使用参数模式
    if args.input_file:
        file_path = args.input_file

        # 验证输入文件
        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return

        if not file_path.lower().endswith('.docx'):
            print("❌ 只支持 .docx 格式的Word文档")
            return

        # 确定输出路径
        if args.overwrite:
            output_path = file_path
        elif args.output:
            output_path = args.output
            if not output_path.lower().endswith('.docx'):
                output_path += '.docx'
        else:
            # 默认保存到当前目录
            base_name = os.path.basename(file_path)
            name, ext = os.path.splitext(base_name)
            output_path = os.path.join(os.getcwd(), f"{name}_processed{ext}")

    else:
        # 交互模式
        # 2. 获取输入文件
        while True:
            file_path = input("\n📂 请输入Word文档路径: ").strip().strip('"')
            if not file_path:
                print("❌ 文件路径不能为空")
                continue

            if not os.path.exists(file_path):
                print(f"❌ 文件不存在: {file_path}")
                continue

            if not file_path.lower().endswith('.docx'):
                print("❌ 只支持 .docx 格式的Word文档")
                continue

            break

    # 3. 检查$符号
    if not check_dollar_signs(file_path):
        return

    # 4. LaTeX拼写检查
    print("\n" + "="*70)
    print("🔤 LaTeX拼写检查")
    print("="*70)
    
    spell_check_passed = check_document_latex_spelling(file_path)
    
    if not spell_check_passed:
        print("\n⚠️  检测到LaTeX拼写错误（如命令中包含空格）")
        choice = input("是否继续处理文档？(y/n): ").strip().lower()
        if choice != 'y':
            print("❌ 操作已取消")
            return
        print("⏩ 继续处理...")

    # 5. 选择保存模式
    output_path = get_save_mode(file_path)    # 3/5. 检查$符号（参数模式也需要检查）
    if not check_dollar_signs(file_path):
        return
    
    # 4/6. LaTeX拼写检查
    print("\n" + "="*70)
    print("🔤 LaTeX拼写检查")
    print("="*70)
    
    spell_check_passed = check_document_latex_spelling(file_path)
    
    if not spell_check_passed:
        print("\n⚠️  检测到LaTeX拼写错误（如命令中包含空格）")
        
        # 在命令行模式下，也要询问是否继续
        if args.input_file:
            choice = input("是否继续处理文档？(y/n): ").strip().lower()
            if choice != 'y':
                print("❌ 操作已取消")
                return
            print("⏩ 继续处理...")
        else:
            choice = input("是否继续处理文档？(y/n): ").strip().lower()
            if choice != 'y':
                print("❌ 操作已取消")
                return
            print("⏩ 继续处理...")
    
    # 检查并关闭Word文档（如果已打开）
    print("\n" + "="*70)
    print("🔍 检查文档状态...")
    print("="*70)
    if not check_and_close_word_document(file_path):
        print("❌ 无法继续：文档仍在使用中")
        return

    # 7/8. 确认操作（仅交互模式）
    if not args.input_file:
        print(f"\n🔄 准备执行:")
        print(f"   输入文件: {file_path}")
        print(f"   输出文件: {output_path}")

        if output_path == file_path:
            confirm = input("\n⚠️  将覆盖原文件，确定继续？(y/n): ").strip().lower()
            if confirm != 'y':
                print("❌ 操作已取消")
                return

    # 8/9. 执行处理
    success = process_document(file_path, output_path)

    if success:
        print(f"\n🎉 文档准备完成！")
        print(f"💡 提示: 已生成 {output_path}")
        
        # 询问是否自动转换和渲染
        print("\n" + "="*70)
        print("🎨 自动转换与渲染")
        print("="*70)
        print("现在将使用Word原生功能自动处理公式：")
        print("   第一步：将 $LaTeX$ 标记为公式对象（模拟Alt+=）")
        print("   第二步：渲染为专业的二维格式")
        
        # 在命令行模式下，也询问是否渲染
        render_choice = input("\n是否自动转换并渲染公式？(y/n): ").strip().lower()
        
        if render_choice == 'y':
            render_success = auto_render_formulas(output_path)
            
            if render_success:
                print("\n✅ 完整流程已完成！")
                print("📂 Word文档已打开，您可以查看渲染后的公式效果")
            else:
                print("\n⚠️  自动渲染未成功，但文档已生成")
                print("💡 您可以手动打开文档并在Word中操作")
        else:
            print("\n⏭️  跳过自动渲染")
            print("💡 您可以稍后手动打开文档")
            
            # 询问是否打开文件
            open_choice = input("\n是否现在打开文档？(y/n): ").strip().lower()
            if open_choice == 'y':
                try:
                    os.startfile(output_path)
                    print(f"✅ 已打开文档")
                except:
                    print(f"❌ 无法自动打开，请手动打开: {output_path}")
    else:
        print(f"\n❌ 处理失败！")

    # 仅在交互模式下等待退出
    if not args.input_file:
        input("\n按Enter键退出...")

if __name__ == "__main__":
    main()
